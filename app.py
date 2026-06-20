"""
Trip Logger Pro — Streamlit App
เก็บข้อมูลถาวรบน Google Sheets (กันข้อมูลหายเวลา redeploy/restart)
เก็บรูปภาพบน Google Drive (เก็บแค่ URL ไว้ใน Sheet ไม่เก็บไฟล์ดิบ)
ส่งอีเมลแจ้งเตือนอัตโนมัติทุกครั้งที่กดบันทึกข้อมูล (โรงแรม/น้ำมัน/รถ)

โครงสร้างข้อมูลใช้ 3 ชีต (worksheet/tab) แยกตามประเภทข้อมูล — แบบ relational
ไม่ใช่เก็บเป็น JSON ก้อนเดียว เพื่อให้เปิดดู/แก้ไขข้อมูลตรงๆ ใน Google Sheets ได้ง่าย:

  ชีต "sites"  — 1 แถวต่อไซต์งาน (ข้อมูลรถ/ไมล์)
    คอลัมน์: SiteName, CreatedAt, UpdatedAt, StartMile, EndMile,
             StartImg_URL, EndImg_URL, CarImg_URL

  ชีต "hotels" — 1 แถวต่อ "ช่องโรงแรม" หนึ่งช่อง (เชื่อมกับ sites ด้วย SiteName)
    คอลัมน์: SiteName, HotelNo, ItemNo, Desc, Locked, Img_URL

  ชีต "fuel"   — 1 แถวต่อ "การเติมน้ำมัน" หนึ่งครั้ง (เชื่อมกับ sites ด้วย SiteName)
    คอลัมน์: SiteName, FuelNo, Date, Province, Locked, Bill_URL, Pre_URL, Post_URL

ถ้า worksheet "hotels"/"fuel" ยังไม่มีในไฟล์ Google Sheet ของคุณ ระบบจะสร้าง
ให้อัตโนมัติพร้อมใส่ header แถวแรกให้เองตอนกดบันทึกครั้งแรก ไม่ต้องสร้างมือ

═══════════════════════════════════════════════════════════════════════════
การตั้งค่าที่ต้องทำก่อนใช้งาน (1 ครั้ง)
═══════════════════════════════════════════════════════════════════════════
1. ใน Google Sheet ที่มีอยู่แล้ว (ชีต "sites") แชร์สิทธิ์ "Editor"
   ให้กับอีเมลของ Service Account (อยู่ในไฟล์ JSON ที่ key "client_email")
   ทำเช่นเดียวกันกับโฟลเดอร์ Google Drive ที่จะใช้เก็บรูป (แชร์ "Editor")

2. ใน Streamlit Cloud → Settings → Secrets ใส่ค่าต่อไปนี้ (รูปแบบ TOML):

   [gcp_service_account]
   type = "service_account"
   project_id = "..."
   private_key_id = "..."
   private_key = "-----BEGIN PRIVATE KEY-----\n...\n-----END PRIVATE KEY-----\n"
   client_email = "...@....iam.gserviceaccount.com"
   client_id = "..."
   auth_uri = "https://accounts.google.com/o/oauth2/auth"
   token_uri = "https://oauth2.googleapis.com/token"
   auth_provider_x509_cert_url = "https://www.googleapis.com/oauth2/v1/certs"
   client_x509_cert_url = "..."

   [app]
   sheet_id = "116BUUyoaU28RMcMWw_-MRpvbuQbbgWUNmRUNm1699I4"   # จาก URL ของคุณ
   drive_folder_id = "เลขที่อยู่ใน URL ของโฟลเดอร์ Google Drive ที่จะเก็บรูป"

   [email]
   sender = "your-email@gmail.com"
   app_password = "App Password 16 หลัก (ไม่ใช่รหัสผ่าน Gmail ปกติ)"
   recipient = "sawitreephi@cpall.co.th"

   ถ้ารันบนเครื่องตัวเอง (local) ให้สร้างไฟล์ .streamlit/secrets.toml
   ใส่เนื้อหาแบบเดียวกันแทน

3. ติดตั้ง dependencies เพิ่มจากเดิม (ใส่ใน requirements.txt):
   gspread
   google-auth
   google-auth-oauthlib
   google-api-python-client
═══════════════════════════════════════════════════════════════════════════
"""

import streamlit as st
import io
import os
import datetime
import json
import smtplib
import base64
import traceback
from email.message import EmailMessage
from docx import Document
from docx.shared import Inches
import streamlit_authenticator as stauth
import yaml
from yaml.loader import SafeLoader
from PIL import Image

import gspread
from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload
import gspread
from oauth2client.service_account import ServiceAccountCredentials

def get_sheet():
    # 1. กำหนด Scope ให้เข้าถึง Sheets และ Drive
    scope = ['https://spreadsheets.google.com/feeds', 'https://www.googleapis.com/auth/drive']
    
    # 2. เชื่อมต่อด้วยไฟล์ JSON (ต้องวางไฟล์นี้ไว้ในโฟลเดอร์เดียวกับโค้ด)
    creds = ServiceAccountCredentials.from_json_keyfile_name('service_account.json', scope)
    client = gspread.authorize(creds)
    
    # 3. เชื่อมต่อด้วย Sheet ID
    SHEET_ID = "116BUUyoaU28RMcMWw_-MRpvbuQbbgWUNmRUNm1699I4"
    sheet = client.open_by_key(SHEET_ID).sheet1
    return sheet
SHEET_ID = "116BUUyoaU28RMcMWw_-MRpvbuQbbgWUNmRUNm1699I4"
FOLDER_ID = "1SbEV8l7oKemwgLg-X7hwOkv0KjByF0wg"
# ─────────────────────────────────────────────────────────────────────────────
# 0. ค่าคงที่
# ─────────────────────────────────────────────────────────────────────────────
DATA_FILE  = "users.json"          # fallback local (เมื่อ Sheets ใช้ไม่ได้)
SCOPES = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive",
]
WORKSHEET_SITES  = "sites"          # ชีตที่คุณสร้างไว้แล้ว: 1 แถวต่อไซต์งาน (รถ/ไมล์)
WORKSHEET_HOTELS = "hotels"         # ชีตใหม่: 1 แถวต่อ "ช่องโรงแรม" หนึ่งช่อง เชื่อมด้วย SiteName
WORKSHEET_FUEL   = "fuel"           # ชีตใหม่: 1 แถวต่อ "การเติมน้ำมัน" หนึ่งครั้ง เชื่อมด้วย SiteName
HOTEL_RANGE = range(1, 4)          # โรงแรม 1-3
HOTEL_ITEM_RANGE = range(1, 7)     # แต่ละโรงแรมมี 6 ช่อง
FUEL_RANGE = range(1, 21)          # การเติมน้ำมัน 1-20 ครั้ง

SITES_HEADER  = ["SiteName", "CreatedAt", "UpdatedAt", "StartMile", "EndMile",
                  "StartImg_URL", "EndImg_URL", "CarImg_URL"]
HOTELS_HEADER = ["SiteName", "HotelNo", "ItemNo", "Desc", "Locked", "Img_URL"]
FUEL_HEADER   = ["SiteName", "FuelNo", "Date", "Province", "Locked",
                  "Bill_URL", "Pre_URL", "Post_URL"]


# ─────────────────────────────────────────────────────────────────────────────
# 1. Google API clients (cache ไว้ใช้ซ้ำ ไม่ auth ใหม่ทุกครั้ง)
# ─────────────────────────────────────────────────────────────────────────────
@st.cache_resource(show_spinner=False)
def _get_gcp_credentials():
    """อ่าน Service Account credentials จาก st.secrets"""
    info = dict(st.secrets["gcp_service_account"])
    return Credentials.from_service_account_info(info, scopes=SCOPES)


@st.cache_resource(show_spinner=False)
def _get_gspread_client():
    creds = _get_gcp_credentials()
    return gspread.authorize(creds)


@st.cache_resource(show_spinner=False)
def _get_drive_service():
    creds = _get_gcp_credentials()
    return build("drive", "v3", credentials=creds)


def _get_or_create_worksheet(sh, title: str, header: list):
    """เปิด worksheet ถ้ามีอยู่แล้ว ถ้าไม่มีให้สร้างใหม่พร้อมใส่ header แถวแรก"""
    try:
        ws = sh.worksheet(title)
        # ถ้า worksheet มีอยู่แต่ไม่มี header (แถวว่าง) ให้ใส่ header ให้
        existing = ws.row_values(1)
        if not existing:
            ws.update([header], value_input_option="RAW")
        return ws
    except gspread.WorksheetNotFound:
        ws = sh.add_worksheet(title=title, rows=500, cols=len(header) + 2)
        ws.update([header], value_input_option="RAW")
        return ws


def _get_sheet():
    gc = _get_gspread_client()
    sheet_id = st.secrets["app"]["sheet_id"]
    return gc.open_by_key(sheet_id)


def _get_sites_ws():
    return _get_or_create_worksheet(_get_sheet(), WORKSHEET_SITES, SITES_HEADER)


def _get_hotels_ws():
    return _get_or_create_worksheet(_get_sheet(), WORKSHEET_HOTELS, HOTELS_HEADER)


def _get_fuel_ws():
    return _get_or_create_worksheet(_get_sheet(), WORKSHEET_FUEL, FUEL_HEADER)


# ─────────────────────────────────────────────────────────────────────────────
# 2. Google Sheets persistence — แทนที่การเขียนไฟล์ JSON ลง disk เดิม
#    (ไฟล์บน disk ของ Streamlit Cloud หายทุกครั้งที่ redeploy/restart
#     ทำให้ "บันทึกข้อมูลไปแล้วหาย" — Sheets แก้ปัญหานี้เพราะเก็บถาวรภายนอก)
# ─────────────────────────────────────────────────────────────────────────────
def json_serial(obj):
    """ตัวช่วยแปลง date object → string (เผื่อใช้ในจุดอื่นที่ยัง serialize JSON อยู่)"""
    if isinstance(obj, (datetime.date, datetime.datetime)):
        return obj.isoformat()
    raise TypeError(f"Type {type(obj)} not serializable")


def _bool_to_cell(v) -> str:
    return "TRUE" if v else "FALSE"


def _cell_to_bool(v) -> bool:
    return str(v).strip().upper() in ("TRUE", "1", "YES", "ใช่")


def _empty_site(site_name: str) -> dict:
    """โครงสร้างไซต์งานเปล่า (ใช้ทั้งตอนสร้างใหม่และตอนเติม default ให้ไซต์ที่ขาดบางส่วน)"""
    return {
        "created_at": "",
        "updated_at": "",
        "hotels": {
            h: {i: {"img": None, "img_id": None, "desc": "", "locked": False}
                for i in HOTEL_ITEM_RANGE}
            for h in HOTEL_RANGE
        },
        "fuel": {
            n: {
                "bill": None, "pre": None, "post": None,
                "bill_id": None, "pre_id": None, "post_id": None,
                "locked": False,
                "date": datetime.date.today(),
                "province": "",
            }
            for n in FUEL_RANGE
        },
        "start_mile": 0,
        "end_mile": 0,
        "start_img": None, "end_img": None, "car_img": None,
    }


def load_sites() -> dict:
    """
    โหลดข้อมูลทุกไซต์จาก Google Sheets — รวม 3 ชีต (sites/hotels/fuel) เข้าด้วยกัน
    เป็นโครงสร้าง dict เดียวกับที่ UI ใช้งาน (st.session_state.sites)
    """
    try:
        sites: dict = {}

        # ── 1. ชีต sites: ข้อมูลระดับไซต์งาน (รถ/ไมล์) ──
        ws_sites = _get_sites_ws()
        for row in ws_sites.get_all_records():
            name = str(row.get("SiteName", "")).strip()
            if not name:
                continue
            site = _empty_site(name)
            site["created_at"] = row.get("CreatedAt", "")
            site["updated_at"] = row.get("UpdatedAt", "")
            try:
                site["start_mile"] = int(row.get("StartMile") or 0)
            except (ValueError, TypeError):
                site["start_mile"] = 0
            try:
                site["end_mile"] = int(row.get("EndMile") or 0)
            except (ValueError, TypeError):
                site["end_mile"] = 0
            site["start_img"] = row.get("StartImg_URL") or None
            site["end_img"]   = row.get("EndImg_URL") or None
            site["car_img"]   = row.get("CarImg_URL") or None
            sites[name] = site

        # ── 2. ชีต hotels: เติมข้อมูลโรงแรมเข้าไปในแต่ละไซต์ ──
        ws_hotels = _get_hotels_ws()
        for row in ws_hotels.get_all_records():
            name = str(row.get("SiteName", "")).strip()
            if name not in sites:
                continue  # โรงแรมของไซต์ที่ไม่มีในชีต sites แล้ว (ถูกลบไปแล้ว) ข้ามทิ้ง
            try:
                h = int(row.get("HotelNo"))
                i = int(row.get("ItemNo"))
            except (ValueError, TypeError):
                continue
            if h not in HOTEL_RANGE or i not in HOTEL_ITEM_RANGE:
                continue
            sites[name]["hotels"][h][i] = {
                "desc":   row.get("Desc", ""),
                "locked": _cell_to_bool(row.get("Locked")),
                "img":    row.get("Img_URL") or None,
                "img_id": None,
            }

        # ── 3. ชีต fuel: เติมข้อมูลการเติมน้ำมันเข้าไปในแต่ละไซต์ ──
        ws_fuel = _get_fuel_ws()
        for row in ws_fuel.get_all_records():
            name = str(row.get("SiteName", "")).strip()
            if name not in sites:
                continue
            try:
                n = int(row.get("FuelNo"))
            except (ValueError, TypeError):
                continue
            if n not in FUEL_RANGE:
                continue
            date_val = row.get("Date", "")
            try:
                date_obj = (datetime.date.fromisoformat(str(date_val))
                            if date_val else datetime.date.today())
            except Exception:
                date_obj = datetime.date.today()
            sites[name]["fuel"][n] = {
                "date":     date_obj,
                "province": row.get("Province", ""),
                "locked":   _cell_to_bool(row.get("Locked")),
                "bill":     row.get("Bill_URL") or None,
                "pre":      row.get("Pre_URL") or None,
                "post":     row.get("Post_URL") or None,
                "bill_id":  None, "pre_id": None, "post_id": None,
            }

        return sites
    except Exception as e:
        st.warning(f"⚠️ โหลดข้อมูลจาก Google Sheets ไม่สำเร็จ: {e}\n\n"
                   f"จะเริ่มต้นด้วยข้อมูลว่าง กรุณาตรวจสอบการตั้งค่า Secrets")
        return {}


def save_sites(sites: dict):
    """
    บันทึกข้อมูลทุกไซต์ลง Google Sheets — เขียนทับทั้ง 3 ชีต (sites/hotels/fuel)
    ตามแถวที่มีข้อมูลจริงเท่านั้น (โรงแรม/น้ำมันช่องที่ยังว่างไม่ถูกเขียนแถว
    เพื่อไม่ให้ชีตยาวเกินจำเป็น — โหลดกลับมาจะเติม default ให้เองผ่าน _empty_site)
    """
    try:
        # ── 1. ชีต sites ──
        ws_sites = _get_sites_ws()
        sites_rows = [SITES_HEADER]
        for name, data in sites.items():
            sites_rows.append([
                name,
                data.get("created_at", ""),
                data.get("updated_at", ""),
                data.get("start_mile", 0),
                data.get("end_mile", 0),
                data.get("start_img") or "",
                data.get("end_img") or "",
                data.get("car_img") or "",
            ])
        ws_sites.clear()
        ws_sites.update(sites_rows, value_input_option="RAW")

        # ── 2. ชีต hotels — เขียนเฉพาะช่องที่มีการบันทึก (locked) หรือมีคำอธิบาย ──
        ws_hotels = _get_hotels_ws()
        hotels_rows = [HOTELS_HEADER]
        for name, data in sites.items():
            for h, h_data in data.get("hotels", {}).items():
                for i, item in h_data.items():
                    if not item.get("locked") and not item.get("desc") and not item.get("img"):
                        continue  # ช่องว่างเปล่า ไม่ต้องเขียนแถว
                    hotels_rows.append([
                        name, h, i,
                        item.get("desc", ""),
                        _bool_to_cell(item.get("locked", False)),
                        item.get("img") or "",
                    ])
        ws_hotels.clear()
        ws_hotels.update(hotels_rows, value_input_option="RAW")

        # ── 3. ชีต fuel — เขียนเฉพาะรายการที่มีการบันทึก (locked) หรือมีข้อมูล ──
        ws_fuel = _get_fuel_ws()
        fuel_rows = [FUEL_HEADER]
        for name, data in sites.items():
            for n, item in data.get("fuel", {}).items():
                has_data = (item.get("locked") or item.get("province")
                            or item.get("bill") or item.get("pre") or item.get("post"))
                if not has_data:
                    continue
                date_val = item.get("date")
                date_str = date_val.isoformat() if isinstance(date_val, (datetime.date, datetime.datetime)) else str(date_val or "")
                fuel_rows.append([
                    name, n, date_str,
                    item.get("province", ""),
                    _bool_to_cell(item.get("locked", False)),
                    item.get("bill") or "",
                    item.get("pre") or "",
                    item.get("post") or "",
                ])
        ws_fuel.clear()
        ws_fuel.update(fuel_rows, value_input_option="RAW")

    except Exception as e:
        st.error(f"❌ บันทึกข้อมูลลง Google Sheets ไม่สำเร็จ: {e}")
        with st.expander("รายละเอียด error (สำหรับแก้ปัญหา)"):
            st.code(traceback.format_exc())


# ─────────────────────────────────────────────────────────────────────────────
# 3. Google Drive — อัปโหลดรูปภาพ คืน URL สาธารณะแบบดูได้ (anyone with link)
# ─────────────────────────────────────────────────────────────────────────────
def upload_to_drive(file_bytes: bytes, filename: str, mime_type: str = "image/png"):
    """
    อัปโหลดรูปไป Google Drive โฟลเดอร์ที่ตั้งค่าไว้
    คืนค่า: (view_url, file_id) หรือ (None, None) ถ้าอัปโหลดไม่สำเร็จ
    """
    if not file_bytes:
        return None, None
    try:
        service = _get_drive_service()
        folder_id = st.secrets["app"]["drive_folder_id"]

        media = MediaIoBaseUpload(io.BytesIO(file_bytes), mimetype=mime_type, resumable=False)
        meta = {"name": filename, "parents": [folder_id]}
        created = service.files().create(body=meta, media_body=media, fields="id").execute()
        file_id = created["id"]

        # เปิดสิทธิ์อ่านสาธารณะ (anyone with the link) เพื่อให้ดูรูปได้จาก URL ตรงๆ
        service.permissions().create(
            fileId=file_id, body={"type": "anyone", "role": "reader"}
        ).execute()

        view_url = f"https://drive.google.com/uc?export=view&id={file_id}"
        return view_url, file_id
    except Exception as e:
        st.error(f"❌ อัปโหลดรูปไป Google Drive ไม่สำเร็จ: {e}")
        return None, None


def _upload_field_if_needed(uploaded_file, label: str):
    """
    รับไฟล์จาก st.file_uploader (UploadedFile) → แปลงเป็น PNG → อัปโหลด Drive
    คืนค่า: (url, file_id) — ถ้า uploaded_file เป็น None คืน (None, None)
    ถ้า uploaded_file เป็น string อยู่แล้ว (= URL เดิมที่เคยอัปโหลดไว้) คืนค่าเดิมไว้เฉยๆ
    """
    if uploaded_file is None:
        return None, None
    if isinstance(uploaded_file, str):
        return uploaded_file, None  # เป็น URL อยู่แล้ว ไม่ต้องอัปโหลดซ้ำ

    raw = _to_bytes(uploaded_file)
    if not raw:
        return None, None
    try:
        img = Image.open(io.BytesIO(raw))
        out = io.BytesIO()
        img.save(out, format="PNG")
        png_bytes = out.getvalue()
    except Exception:
        png_bytes = raw  # ถ้าแปลงไม่ได้ ส่งดิบไปเลย

    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_label = "".join(c for c in label if c.isalnum() or c in ("_", "-"))
    filename = f"{safe_label}_{ts}.png"
    return upload_to_drive(png_bytes, filename, "image/png")


# ─────────────────────────────────────────────────────────────────────────────
# 4. ระบบส่งอีเมลแจ้งเตือน — ส่งทุกครั้งที่กดบันทึก (โรงแรม/น้ำมัน/รถ)
# ─────────────────────────────────────────────────────────────────────────────
def send_notification_email(site_name: str, section: str, detail: str = ""):
    """
    ส่งอีเมลแจ้งเตือนแบบข้อความสั้นๆ (ไม่แนบไฟล์) ทุกครั้งที่มีการบันทึก/อัปเดตข้อมูล
    section: "โรงแรม" / "น้ำมัน" / "ข้อมูลรถ"
    """
    try:
        sender    = st.secrets["email"]["sender"]
        password  = st.secrets["email"]["app_password"]
        recipient = st.secrets["email"].get("recipient", "sawitreephi@cpall.co.th")

        msg = EmailMessage()
        msg["Subject"] = f"[Trip Logger] อัปเดตข้อมูล — {site_name} ({section})"
        msg["From"] = sender
        msg["To"] = recipient
        now_text = datetime.datetime.now().strftime("%d/%m/%Y %H:%M")
        body = (
            f"มีการบันทึก/อัปเดตข้อมูลในระบบ Trip Logger\n\n"
            f"ไซต์งาน : {site_name}\n"
            f"ส่วนที่อัปเดต : {section}\n"
            f"เวลา : {now_text}\n"
        )
        if detail:
            body += f"\nรายละเอียด:\n{detail}\n"
        msg.set_content(body)

        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
            smtp.login(sender, password)
            smtp.send_message(msg)
        return True
    except Exception as e:
        st.warning(f"⚠️ ส่งอีเมลแจ้งเตือนไม่สำเร็จ: {e}")
        return False


def send_email_report(site_name: str, buf: io.BytesIO):
    """ส่งรายงาน Word ฉบับเต็มเป็นไฟล์แนบ (ใช้ในหน้า Export โดย admin)"""
    sender    = st.secrets["email"]["sender"]
    password  = st.secrets["email"]["app_password"]
    recipient = st.secrets["email"].get("recipient", "sawitreephi@cpall.co.th")

    msg = EmailMessage()
    msg["Subject"] = f"Trip Report — {site_name}"
    msg["From"] = sender
    msg["To"] = recipient
    msg.set_content(
        f"ส่งรายงานไซต์งาน: {site_name} "
        f"ในวันที่ {datetime.datetime.now().strftime('%d/%m/%Y')}"
    )
    msg.add_attachment(
        buf.getvalue(),
        maintype="application",
        subtype="octet-stream",
        filename=f"Trip_Report_{site_name}.docx",
    )
    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
        smtp.login(sender, password)
        smtp.send_message(msg)


# ─────────────────────────────────────────────────────────────────────────────
# 5. ผู้ใช้งาน (เก็บ local เหมือนเดิม — ไม่กระทบข้อมูลไซต์งานที่หายไปก่อนหน้า)
# ─────────────────────────────────────────────────────────────────────────────
def load_users():
    if os.path.exists(DATA_FILE):
        with open(DATA_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {"usernames": {"admin": {"name": "Admin", "password": "default_password"}}}


def save_users(users):
    with open(DATA_FILE, "w", encoding="utf-8") as f:
        json.dump(users, f, indent=4, ensure_ascii=False)


# ─────────────────────────────────────────────────────────────────────────────
# 6. Init session state (ห้ามมี UI ก่อนหน้านี้)
# ─────────────────────────────────────────────────────────────────────────────
if "credentials" not in st.session_state:
    st.session_state.credentials = load_users()

if "sites" not in st.session_state:
    with st.spinner("กำลังโหลดข้อมูลจาก Google Sheets..."):
        st.session_state.sites = load_sites()


# ─────────────────────────────────────────────────────────────────────────────
# 7. Admin Panel
# ─────────────────────────────────────────────────────────────────────────────
def admin_panel():
    st.subheader("🛠️ Admin Settings: จัดการผู้ใช้งาน")
    with st.expander("➕ เพิ่มผู้ใช้งานใหม่"):
        new_username = st.text_input("Username")
        new_name     = st.text_input("ชื่อจริง")
        new_password = st.text_input("รหัสผ่าน", type="password")
        if st.button("บันทึกผู้ใช้"):
            if new_username and new_password:
                st.session_state.credentials["usernames"][new_username] = {
                    "name": new_name, "password": new_password
                }
                save_users(st.session_state.credentials)
                st.success(f"บันทึกผู้ใช้ {new_username} เรียบร้อย")
            else:
                st.error("กรุณากรอกข้อมูลให้ครบ")

    st.write("### รายชื่อผู้ใช้")
    for username in list(st.session_state.credentials["usernames"].keys()):
        col1, col2 = st.columns([0.8, 0.2])
        col1.write(f"- {username}")
        if col2.button("ลบ", key=f"del_{username}"):
            del st.session_state.credentials["usernames"][username]
            save_users(st.session_state.credentials)
            st.rerun()

    st.write("---")
    with st.expander("🔌 ตรวจสอบการเชื่อมต่อ Google Sheets / Drive"):
        if st.button("ทดสอบการเชื่อมต่อ"):
            try:
                ws_s = _get_sites_ws()
                ws_h = _get_hotels_ws()
                ws_f = _get_fuel_ws()
                st.success(
                    f"✅ เชื่อม Google Sheets สำเร็จ — "
                    f"sites: {len(ws_s.get_all_values())-1} แถว, "
                    f"hotels: {len(ws_h.get_all_values())-1} แถว, "
                    f"fuel: {len(ws_f.get_all_values())-1} แถว"
                )
            except Exception as e:
                st.error(f"❌ เชื่อม Google Sheets ไม่สำเร็จ: {e}")
            try:
                _get_drive_service()
                st.success("✅ เชื่อม Google Drive สำเร็จ")
            except Exception as e:
                st.error(f"❌ เชื่อม Google Drive ไม่สำเร็จ: {e}")

    st.write("---")
    if st.button("รีเซ็ตสถานะหน้าเพจ"):
        st.session_state.page = "site_selector"
        st.rerun()


# ─────────────────────────────────────────────────────────────────────────────
# 8. ฟังก์ชันจัดการไซต์งาน
# ─────────────────────────────────────────────────────────────────────────────
def now_str() -> str:
    return datetime.datetime.now().strftime("%d/%m/%Y %H:%M")


def init_site(site_name: str):
    """สร้างหรือโหลดข้อมูลไซต์ใน session_state"""
    if site_name not in st.session_state.sites:
        new_site = _empty_site(site_name)
        new_site["created_at"] = now_str()
        new_site["updated_at"] = now_str()
        st.session_state.sites[site_name] = new_site
    st.session_state.current_site = site_name
    st.session_state.hotels = st.session_state.sites[site_name]["hotels"]
    st.session_state.fuel   = st.session_state.sites[site_name]["fuel"]


def touch_site():
    """อัปเดต updated_at + บันทึกลง Google Sheets ทันที"""
    site = st.session_state.get("current_site")
    if site and site in st.session_state.sites:
        st.session_state.sites[site]["updated_at"] = now_str()
        save_sites(st.session_state.sites)


# ─────────────────────────────────────────────────────────────────────────────
# 9. Toggle Lock — อัปโหลดรูปไป Drive ตอนล็อก (กดบันทึก) แล้วเก็บแค่ URL
# ─────────────────────────────────────────────────────────────────────────────
def _to_bytes(file_obj):
    if file_obj is None:
        return None
    if isinstance(file_obj, bytes):
        return file_obj
    if isinstance(file_obj, str):
        return None  # เป็น URL อยู่แล้ว ไม่ใช่ bytes
    try:
        file_obj.seek(0)
        return file_obj.read()
    except Exception:
        return None


def toggle_lock(section, key, sub_key=None):
    if section == "hotel":
        item = st.session_state.hotels[key][sub_key]
        if not item["locked"]:
            with st.spinner("กำลังอัปโหลดรูป..."):
                url, fid = _upload_field_if_needed(item["img"], f"hotel{key}_{sub_key}")
                if url:
                    item["img"] = url
                    item["img_id"] = fid
        item["locked"] = not item["locked"]
        section_label = f"โรงแรม {key} / รายการ {sub_key}"
    else:
        item = st.session_state.fuel[key]
        if not item["locked"]:
            with st.spinner("กำลังอัปโหลดรูป..."):
                url, fid = _upload_field_if_needed(item["bill"], f"fuel{key}_bill")
                if url: item["bill"], item["bill_id"] = url, fid
                url, fid = _upload_field_if_needed(item["pre"], f"fuel{key}_pre")
                if url: item["pre"], item["pre_id"] = url, fid
                url, fid = _upload_field_if_needed(item["post"], f"fuel{key}_post")
                if url: item["post"], item["post_id"] = url, fid
        item["locked"] = not item["locked"]
        section_label = f"น้ำมัน ครั้งที่ {key}"

    touch_site()

    # ส่งอีเมลแจ้งเตือนทุกครั้งที่กดบันทึก (ไม่ส่งตอนกด "แก้ไข" คือ unlock)
    if item["locked"]:
        site = st.session_state.get("current_site", "")
        send_notification_email(site, section_label)


# ─────────────────────────────────────────────────────────────────────────────
# 10. สร้างไฟล์ Word
# ─────────────────────────────────────────────────────────────────────────────
def _img_stream(data):
    """
    คืน BytesIO พร้อมใช้กับ docx
    รองรับทั้ง: bytes ดิบ, UploadedFile, และ URL string (ดาวน์โหลดจาก Drive ก่อน)
    """
    if data is None:
        return None
    try:
        if isinstance(data, str):
            # เป็น URL จาก Drive — ดาวน์โหลดมาก่อน
            import urllib.request
            with urllib.request.urlopen(data, timeout=15) as resp:
                raw = resp.read()
        elif isinstance(data, bytes):
            raw = data
        else:
            data.seek(0)
            raw = data.read()
        if not raw:
            return None

        img = Image.open(io.BytesIO(raw))
        out_stream = io.BytesIO()
        img.save(out_stream, format="PNG")
        out_stream.seek(0)
        return out_stream
    except Exception:
        return None


def _safe_add_picture(run, data, width):
    stream = _img_stream(data)
    if stream is None:
        return False
    try:
        run.add_picture(stream, width=width)
        return True
    except Exception:
        return False


def generate_word():
    doc  = Document()
    site = st.session_state.get("current_site", "")
    info = st.session_state.sites.get(site, {})
    doc.add_heading(f"Trip Report — {site}", 0)
    doc.add_paragraph(f"สร้างเมื่อ: {info.get('created_at', '-')}   อัปเดต: {info.get('updated_at', '-')}")

    # ── ส่วนที่ 1: ข้อมูลรถ ──
    doc.add_heading("ข้อมูลรถและการเดินทาง", level=1)
    start_mile = info.get("start_mile") or st.session_state.get("start_mile") or 0
    end_mile   = info.get("end_mile")   or st.session_state.get("end_mile")   or 0
    distance   = max(0, end_mile - start_mile)
    doc.add_paragraph(f"เลขไมล์เริ่มต้น : {start_mile}")
    doc.add_paragraph(f"เลขไมล์หลังจบ  : {end_mile}")
    doc.add_paragraph(f"ระยะทางรวม     : {distance} กม.")

    car_images = []
    for label, key in [("ไมล์ก่อนเริ่ม", "start_img"),
                        ("ไมล์หลังจบ",    "end_img"),
                        ("รูปรถทั่วไป",   "car_img")]:
        raw = info.get(key) or st.session_state.get(key)
        if raw is not None:
            car_images.append((label, raw))

    if car_images:
        tbl = doc.add_table(rows=2, cols=len(car_images))
        tbl.style = "Table Grid"
        for ci, (label, raw_data) in enumerate(car_images):
            tbl.rows[0].cells[ci].text = label
            p   = tbl.rows[1].cells[ci].paragraphs[0]
            run = p.add_run()
            _safe_add_picture(run, raw_data, Inches(1.8))
    doc.add_paragraph("")

    # ── ส่วนที่ 2: โรงแรม ──
    doc.add_heading("รูปภาพโรงแรม", level=1)
    any_hotel = False
    for h in HOTEL_RANGE:
        items_in_hotel = [
            (i, st.session_state.hotels[h][i])
            for i in HOTEL_ITEM_RANGE
            if st.session_state.hotels[h][i]["locked"]
            and st.session_state.hotels[h][i]["img"] is not None
        ]
        if not items_in_hotel:
            continue
        any_hotel = True
        doc.add_heading(f"โรงแรมที่ {h}", level=2)
        for i, item in items_in_hotel:
            doc.add_paragraph(f"รายการที่ {i}: {item['desc']}")
            run = doc.add_paragraph().add_run()
            _safe_add_picture(run, item["img"], Inches(2.5))
    if not any_hotel:
        doc.add_paragraph("(ยังไม่มีรูปโรงแรมที่บันทึกแล้ว)")

    # ── ส่วนที่ 3: น้ำมัน ──
    fuel_count = 0
    for n in FUEL_RANGE:
        item = st.session_state.fuel.get(n)
        if not item or not item.get("locked"):
            continue

        fuel_count += 1
        doc.add_heading(f"การเติมครั้งที่ {n}", level=2)
        doc.add_paragraph(f"วันที่: {item.get('date', '-')}   จังหวัด: {item.get('province', '-')}")

        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        tbl.autofit = False
        tbl.columns[0].width = Inches(3.5)
        tbl.columns[1].width = Inches(3.5)

        cell_left = tbl.cell(0, 0)
        p_left = cell_left.paragraphs[0]
        if item.get("bill"):
            _safe_add_picture(p_left.add_run(), item["bill"], Inches(4.2))
        else:
            p_left.text = "(ไม่มีรูปใบเสร็จ)"

        cell_right = tbl.cell(0, 1)
        p_pre = cell_right.add_paragraph("ไมล์ก่อนเติม:")
        if item.get("pre"):
            _safe_add_picture(p_pre.add_run(), item["pre"], Inches(2.0))
        p_post = cell_right.add_paragraph("\nไมล์หลังเติม:")
        if item.get("post"):
            _safe_add_picture(p_post.add_run(), item["post"], Inches(2.0))

        doc.add_paragraph("\n")

    if fuel_count == 0:
        doc.add_paragraph("(ยังไม่มีรายการเติมน้ำมันที่บันทึกแล้ว)")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─────────────────────────────────────────────────────────────────────────────
# 11. UI: โรงแรม
# ─────────────────────────────────────────────────────────────────────────────
def render_hotel_section():
    st.header("🏨 ส่วนที่ 1: รูปภาพโรงแรม")
    for h in HOTEL_RANGE:
        with st.expander(f"โรงแรมที่ {h}"):
            cols = st.columns(3)
            for i in HOTEL_ITEM_RANGE:
                item = st.session_state.hotels[h][i]
                with cols[(i - 1) % 3]:
                    if not item["locked"]:
                        up = st.file_uploader("เลือกรูป", key=f"up_h_{h}_{i}")
                        if up:
                            item["img"] = up
                        item["desc"] = st.text_input(
                            "คำอธิบาย", value=item["desc"], key=f"desc_h_{h}_{i}"
                        )
                        if st.button("บันทึก", key=f"btn_h_save_{h}_{i}"):
                            toggle_lock("hotel", h, i)
                            st.rerun()
                    else:
                        if item["img"] is not None:
                            st.image(item["img"], use_container_width=True)
                        st.info(f"📝 {item['desc']}")
                        if st.button("แก้ไข", key=f"btn_h_edit_{h}_{i}"):
                            toggle_lock("hotel", h, i)
                            st.rerun()


# ─────────────────────────────────────────────────────────────────────────────
# 12. UI: ข้อมูลรถ
# ─────────────────────────────────────────────────────────────────────────────
def render_car_section():
    st.subheader("📋 ข้อมูลรถและการเดินทาง")
    site     = st.session_state.get("current_site", "")
    site_inf = st.session_state.sites.get(site, {})

    with st.container(border=True):
        col_a, col_b = st.columns(2)
        with col_a:
            st.file_uploader("รูปไมล์ก่อนเริ่ม", key="start_img")
            st.number_input("เลขไมล์เริ่มต้น", key="start_mile", min_value=0)
            show = st.session_state.get("start_img") or site_inf.get("start_img")
            if show:
                st.image(show, caption="ไมล์ก่อนเริ่ม ✅" if site_inf.get("start_img") else "ไมล์ก่อนเริ่ม")
        with col_b:
            st.file_uploader("รูปไมล์หลังจบทริป", key="end_img")
            st.number_input("เลขไมล์หลังจบ", key="end_mile", min_value=0)
            show = st.session_state.get("end_img") or site_inf.get("end_img")
            if show:
                st.image(show, caption="ไมล์หลังจบ ✅" if site_inf.get("end_img") else "ไมล์หลังจบ")

    st.file_uploader("อัปโหลดรูปรถ (ทั่วไป)", key="car_img")
    show = st.session_state.get("car_img") or site_inf.get("car_img")
    if show:
        st.image(show, caption="รูปรถ ✅" if site_inf.get("car_img") else "รูปรถ")

    if st.button("💾 บันทึกข้อมูลรถ"):
        site = st.session_state.get("current_site")
        if site:
            with st.spinner("กำลังอัปโหลดรูปและบันทึก..."):
                s = st.session_state.sites[site]
                s["start_mile"] = st.session_state.get("start_mile", 0)
                s["end_mile"]   = st.session_state.get("end_mile", 0)
                for key in ["start_img", "end_img", "car_img"]:
                    uploaded = st.session_state.get(key)
                    if uploaded is not None and not isinstance(uploaded, str):
                        url, _fid = _upload_field_if_needed(uploaded, f"{site}_{key}")
                        if url:
                            s[key] = url
                touch_site()
            send_notification_email(site, "ข้อมูลรถ",
                                     f"เลขไมล์เริ่มต้น {s['start_mile']} → "
                                     f"เลขไมล์หลังจบ {s['end_mile']}")
        st.success("✅ บันทึกข้อมูลรถเรียบร้อย")


# ─────────────────────────────────────────────────────────────────────────────
# 13. UI: น้ำมัน
# ─────────────────────────────────────────────────────────────────────────────
def render_fuel_section():
    st.header("⛽ ส่วนที่ 2: บันทึกการเติมน้ำมัน")
    render_car_section()
    st.divider()

    for n in FUEL_RANGE:
        with st.expander(f"การเติมครั้งที่ {n}"):
            item = st.session_state.fuel[n]

            if not item["locked"]:
                c1, c2, c3 = st.columns(3)
                with c1:
                    up_bill = st.file_uploader("📸 ใบเสร็จ", key=f"up_bill_{n}")
                    if up_bill:
                        item["bill"] = up_bill
                with c2:
                    up_pre = st.file_uploader("🔼 ไมล์ก่อนเติม", key=f"pre_{n}")
                    if up_pre:
                        item["pre"] = up_pre
                with c3:
                    up_post = st.file_uploader("🔽 ไมล์หลังเติม", key=f"post_{n}")
                    if up_post:
                        item["post"] = up_post

                c_date, c_prov = st.columns(2)
                with c_date:
                    item["date"] = st.date_input(
                        "วันที่เติม",
                        value=item.get("date") or datetime.date.today(),
                        key=f"date_{n}",
                    )
                with c_prov:
                    item["province"] = st.text_input(
                        "จังหวัด", value=item.get("province", ""), key=f"prov_{n}"
                    )

                if st.button("💾 บันทึกรายการ", key=f"save_f_{n}"):
                    toggle_lock("fuel", n)
                    st.rerun()
            else:
                st.write("✅ บันทึกรายการนี้แล้ว")
                st.info(f"📅 วันที่: {item['date']} | 📍 จังหวัด: {item['province']}")
                col_receipt, col_miles = st.columns([3, 1])

                with col_receipt:
                    if item.get("bill"):
                        st.image(item["bill"], caption="ใบเสร็จ", use_container_width=True)

                with col_miles:
                    if item.get("pre"):
                        st.image(item["pre"], caption="ไมล์ก่อนเติม", width=250)
                    if item.get("post"):
                        st.image(item["post"], caption="ไมล์หลังเติม", width=250)

                if st.button("✏️ แก้ไขรายการ", key=f"edit_f_{n}"):
                    toggle_lock("fuel", n)
                    st.rerun()


# ─────────────────────────────────────────────────────────────────────────────
# 14. UI: Export
# ─────────────────────────────────────────────────────────────────────────────
def render_export_section():
    site = st.session_state.get("current_site", "trip")
    filename = f"Trip_Report_{site}.docx"
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    ok = False
    data = None
    err = ""
    try:
        buf = generate_word()
        data = buf.getvalue()
        ok = True
    except Exception as e:
        ok = False
        err = str(e)
        st.error(f"สร้างไฟล์ไม่ได้: {e}")

    if ok:
        st.divider()
        st.header("📤 Export และส่งรายงาน")

        st.download_button(
            label="📥 ดาวน์โหลด Word",
            data=data,
            file_name=filename,
            mime=mime,
            use_container_width=True,
        )

        if st.session_state.get("username") == "admin":
            if st.button("📧 ส่งรายงานผ่านอีเมล (Admin Only)", type="secondary", use_container_width=True):
                try:
                    send_email_report(site, buf)
                    st.success("✅ ส่งอีเมลรายงานเรียบร้อยแล้ว")
                except Exception as e:
                    st.error(f"❌ ส่งอีเมลไม่สำเร็จ: {e}")

    hotel_count = sum(
        1 for h in HOTEL_RANGE for i in HOTEL_ITEM_RANGE
        if st.session_state.hotels[h][i]["locked"]
        and st.session_state.hotels[h][i]["img"] is not None
    )
    fuel_count = sum(
        1 for n in FUEL_RANGE
        if st.session_state.fuel[n]["locked"]
        and st.session_state.fuel[n]["bill"] is not None
    )

    col_dl, col_info = st.columns([1, 2])
    with col_info:
        st.info(
            f"📦 รายงานประกอบด้วย\n"
            f"- 🏨 รูปโรงแรมที่บันทึกแล้ว: **{hotel_count}** รายการ\n"
            f"- ⛽ การเติมน้ำมันที่บันทึกแล้ว: **{fuel_count}** ครั้ง"
        )
    with col_dl:
        if ok:
            st.download_button(
                label="📥 ดาวน์โหลดรายงาน Word",
                data=data, file_name=filename, mime=mime,
                key="dl_main", type="primary", use_container_width=True,
            )
        else:
            st.error(f"❌ สร้างไฟล์ Word ไม่ได้\n\n`{err}`")


# ─────────────────────────────────────────────────────────────────────────────
# 15. UI: หน้าเลือก/สร้างไซต์งาน
# ─────────────────────────────────────────────────────────────────────────────
def render_site_selector():
    st.title("🏗️ ไซต์งาน")

    sites = st.session_state.sites
    if sites:
        st.subheader("📋 ไซต์งานที่บันทึกไว้")
        for site_name, info in sorted(
            sites.items(),
            key=lambda x: x[1].get("updated_at", ""),
            reverse=True,
        ):
            with st.container(border=True):
                col_name, col_meta, col_btn = st.columns([2, 3, 1])

                with col_name:
                    st.markdown(f"### 📁 {site_name}")

                with col_meta:
                    created = info.get("created_at", "-")
                    updated = info.get("updated_at", "-")
                    h_count = sum(
                        1
                        for h_data in info.get("hotels", {}).values()
                        for item in h_data.values()
                        if item.get("locked") and item.get("img") is not None
                    )
                    f_count = sum(
                        1
                        for item in info.get("fuel", {}).values()
                        if item.get("locked") and item.get("bill") is not None
                    )
                    st.caption(f"🕐 สร้าง: {created}")
                    st.caption(f"🔄 อัปเดต: {updated}")
                    st.caption(f"🏨 โรงแรม: {h_count} รายการ  |  ⛽ น้ำมัน: {f_count} ครั้ง")

                with col_btn:
                    if st.button("เข้าแก้ไข", key=f"open_{site_name}", use_container_width=True, type="primary"):
                        init_site(site_name)
                        st.session_state.page = "trip_logger"
                        st.rerun()
                    if st.button("🗑️ ลบ", key=f"del_site_{site_name}", use_container_width=True):
                        del st.session_state.sites[site_name]
                        save_sites(st.session_state.sites)
                        st.rerun()
    else:
        st.info("ยังไม่มีไซต์งาน กรุณาสร้างไซต์งานใหม่ด้านล่าง")

    st.divider()

    st.subheader("➕ สร้างไซต์งานใหม่")
    col_input, col_btn2 = st.columns([3, 1])
    with col_input:
        site_name = st.text_input("ชื่อไซต์งาน", label_visibility="collapsed", placeholder="ระบุชื่อไซต์งาน...")
    with col_btn2:
        if st.button("สร้าง", type="primary", use_container_width=True):
            if site_name:
                if site_name in st.session_state.sites:
                    st.warning(f"ไซต์งาน '{site_name}' มีอยู่แล้ว กดเข้าแก้ไขได้เลย")
                else:
                    init_site(site_name)
                    save_sites(st.session_state.sites)
                    st.session_state.page = "trip_logger"
                    st.rerun()
            else:
                st.warning("กรุณาระบุชื่อไซต์งาน")


# ─────────────────────────────────────────────────────────────────────────────
# 16. Main App
# ─────────────────────────────────────────────────────────────────────────────
st.set_page_config(layout="wide", page_title="Trip Logger Pro")

with open("config.yaml", encoding="utf-8") as file:
    config = yaml.load(file, Loader=SafeLoader)

authenticator = stauth.Authenticate(
    config["credentials"],
    config["cookie"]["name"],
    config["cookie"]["key"],
    config["cookie"]["expiry_days"],
)

authenticator.login()

if st.session_state["authentication_status"] is False:
    st.error("ชื่อผู้ใช้หรือรหัสผ่านไม่ถูกต้อง")

elif st.session_state["authentication_status"] is None:
    st.warning("กรุณากรอกชื่อผู้ใช้และรหัสผ่านเพื่อเข้าสู่ระบบ")

elif st.session_state["authentication_status"]:
    st.sidebar.write(f'ยินดีต้อนรับ *{st.session_state["name"]}*')
    authenticator.logout("Logout", "sidebar")

    if st.session_state.get("username") == "admin":
        with st.sidebar.expander("Admin Settings"):
            admin_panel()

    if "page" not in st.session_state:
        st.session_state.page = "site_selector"

    if st.session_state.page == "site_selector":
        render_site_selector()

    elif st.session_state.page == "trip_logger":
        site = st.session_state.get("current_site", "")

        if "hotels" not in st.session_state or "fuel" not in st.session_state:
            init_site(site)

        st.sidebar.info(f"📁 ไซต์งาน: {site}")
        if st.sidebar.button("⬅️ กลับหน้าไซต์งาน"):
            save_sites(st.session_state.sites)
            st.session_state.page = "site_selector"
            st.rerun()

        st.title(f"📸 Trip Logger: {site}")

        render_hotel_section()
        st.divider()
        render_fuel_section()
        render_export_section()
